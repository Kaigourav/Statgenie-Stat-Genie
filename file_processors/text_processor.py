import pandas as pd
from docx import Document

def read_txt(path: str) -> pd.DataFrame:
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        lines = [line.strip() for line in f.readlines() if line.strip()]
    return pd.DataFrame({"text": lines})

def read_docx(path: str) -> pd.DataFrame:
    doc = Document(path)
    tables = []
    for tbl in doc.tables:
        data = [[cell.text.strip() for cell in row.cells] for row in tbl.rows]
        if data:
            tables.append(pd.DataFrame(data[1:], columns=data[0]))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    
    if tables:
        df_tables = pd.concat(tables, ignore_index=True)
        df_tables['__doc_paragraphs'] = ' | '.join(paragraphs)
        return df_tables
    return pd.DataFrame({'text': paragraphs})