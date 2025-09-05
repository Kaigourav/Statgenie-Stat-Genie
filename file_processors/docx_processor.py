import pandas as pd
from docx import Document
from .unstructured_parser import ai_extract_table_to_csv

def read_docx(path: str, csv_output_path: str = None) -> pd.DataFrame:
    """
    Extract tables from DOCX, fallback to AI for paragraphs.
    """
    doc = Document(path)
    tables = []
    for t in doc.tables:
        data = []
        for row in t.rows:
            data.append([cell.text.strip() for cell in row.cells])
        if data:
            tables.append(pd.DataFrame(data[1:], columns=data[0]))

    if tables:
        df = pd.concat(tables, ignore_index=True)
        if csv_output_path:
            df.to_csv(csv_output_path, index=False)
        return df

    # Fallback: extract paragraphs
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    if paragraphs:
        if csv_output_path:
            return ai_extract_table_to_csv(paragraphs, csv_output_path)
        return ai_extract_table_to_csv(paragraphs, None)

    return pd.DataFrame()
